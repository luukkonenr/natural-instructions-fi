
import json
import glob 
import os
import docx
import tqdm 
import argparse
import re

def yield_docs(paths):
    for path in sorted(paths):
        with open(path) as f:
            doc = json.load(f)

        if  doc["Input_language"][0] == doc["Output_language"][0] == doc["Instruction_language"][0] == "English":
            yield (doc, path)

def get_categ_levels(instances):
    categories = set()
    for inst in instances:
        categories.update(inst["output"])

    return list(categories)


if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Take a list of natural instruction task files and convert them into intermediate .docx-format")
    ap.add_argument("files", nargs="+", help="files to translate")
    ap.add_argument("--prefix", default="", help="Add identifier to output files")
    args = ap.parse_args()

    doclist = list()
    doc_categs = dict()
    doc_counter=0
    doc=docx.Document()
    current_len=0
    max_len=970_000
    
    for d_idx, (d, doc_name) in tqdm.tqdm(enumerate(yield_docs(args.files))):
        print(doc_name)
        instances, definition = [d.pop(key) for key in ["Instances", "Definition"]]
        
        # Add task separator and task definition
        pgraph=doc.add_paragraph("")
        header = f"Task number {d_idx}"
        current_len += len(header)
        r = pgraph.add_run(header) # Definition -> task
        r.bold=True
        r.underline=True

        doc.add_paragraph(definition)
        current_len+=len(definition[0])

        # Add category levels with header
        levels = get_categ_levels(instances)
        doc_categs.update({doc_name: levels})
        for level in levels:
            pgraph=doc.add_paragraph("")
            header = f"Level"
            r = pgraph.add_run(header)
            r.underline=True
            current_len+=len(header)
            pgraph=doc.add_paragraph("")
            r = pgraph.add_run(level)
            current_len+=len(level)


        for sample_idx, sample in enumerate(instances):
            sample_len = len(sample["input"])
            if (current_len + sample_len) > max_len:
                #Save first!
                doc.save(f"doc_in/natural_inst-{args.prefix}-{doc_counter:05d}.docx")
                doc_counter+=1
                current_len=0
                doc=docx.Document()

            doclist.append(
                {"file": doc_name, 
                "idx": sample_idx, 
                "outputs": sample["output"], 
                "uuid": sample["id"], 
                })

            pgraph=doc.add_paragraph("")
            header = f"Example {d_idx}.{sample_idx}"
            current_len+= len(header)
            r=pgraph.add_run(header) # adds to the previous paragraph
            r.bold=True
            pattern = re.compile(r'[\x01-\x1F\x7F]') # 
            s_input = pattern.sub("", sample["input"])
            pgraph=doc.add_paragraph(fr'{s_input}')

            current_len+=sample_len

        doc.save(f"doc_in/natural_inst-{args.prefix}-{doc_counter:05d}.docx")


    with open(f"doc_in/metadata-{args.prefix}.json","wt") as f:
        json.dump(doclist,f, indent=2)

    with open(f"doc_in/metadata-{args.prefix}_category_map.json","wt") as f:
        json.dump(doc_categs,f, indent=2)
