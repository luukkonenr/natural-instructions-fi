
import json
import glob 
import os
import docx
import tqdm 
import argparse
import re
# Yield only if task is monolingual

def yield_docs(paths):
    for path in sorted(paths):
        with open(path) as f:
            doc = json.load(f)

        if  doc["Input_language"][0] == doc["Output_language"][0] == doc["Instruction_language"][0] == "English":
            yield (doc, path)



if __name__ == "__main__":
    ap = argparse.ArgumentParser(description="Take a list of natural instruction task files and convert them into intermediate .docx-format")
    ap.add_argument("files", nargs="+", help="files to translate")
    ap.add_argument("--output_prefix", default="", help="Add identifier to output files")
    args = ap.parse_args()

    doclist = list()
    doc_counter=0
    doc=docx.Document()
    current_len=0
    max_len=950_000
    
    for d_idx, (d, doc_name) in tqdm.tqdm(enumerate(yield_docs(args.files))):
        print(doc_name)
        instances, definition = [d.pop(key) for key in ["Instances", "Definition"]]

        pgraph=doc.add_paragraph("")
        r = pgraph.add_run(f"Task number {d_idx}") # Definition -> task
        r.bold=True
        r.underline=True
        doc.add_paragraph(definition)
        current_len+=len(definition[0])

        for sample_idx, sample in enumerate(instances):

            sample_len = len(sample["input"]) + sum([len(o) for o in sample["output"]])
            if (current_len + sample_len) > max_len:
                #Save first!
                doc.save(f"doc_in/natural_inst-{args.output_prefix}-{doc_counter:05d}.docx")
                doc_counter+=1
                current_len=0
                doc=docx.Document()

            doclist.append(
                {"file": doc_name, 
                "idx": sample_idx, 
                "outputs": len(sample["output"]), 
                "uuid": sample["id"], 
                })

            pgraph=doc.add_paragraph("")
            r=pgraph.add_run(f"Example {d_idx}.{sample_idx}") # adds to the previous paragraph
            r.bold=True
            pattern = re.compile(r'[\x01-\x1F\x7F]') # 
            print(sample)
            s_input = pattern.sub("", sample["input"])
            pgraph=doc.add_paragraph(fr'{s_input}')

            for out_idx, entry in enumerate(sample["output"]):
                pgraph=doc.add_paragraph("")
                r = pgraph.add_run(f"Result") # output -> result
                r.bold=True
                doc.add_paragraph(entry)       
    
            current_len+=sample_len

        doc.save(f"doc_in/natural_inst-{args.output_prefix}-{doc_counter:05d}.docx")


    with open(f"doc_in/metadata-{args.output_prefix}.json","wt") as f:
        json.dump(doclist,f)

